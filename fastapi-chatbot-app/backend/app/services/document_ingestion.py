"""Document ingestion service: parses CSV, PDF, DOCX, including images/graphs via OCR,
chunks text, embeds and stores in Qdrant via LangChain.

PDF/DOCX images are OCR'd when possible. If OCR runtime is unavailable (e.g.,
Tesseract not installed), the service falls back to text-only extraction gracefully.
"""

from typing import List, Optional
from pathlib import Path
import io
import pandas as pd
from PyPDF2 import PdfReader
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_qdrant import Qdrant
from app.services.embeddings import GeminiEmbeddings
from app.config import settings

# Optional OCR/vision imports with graceful fallback
import logging

logger = logging.getLogger(__name__)

try:  # Image manipulation
    from PIL import Image
except Exception:  # pragma: no cover - optional
    Image = None  # type: ignore

try:  # OCR binding (requires Tesseract runtime installed separately)
    import pytesseract
except Exception:  # pragma: no cover - optional
    pytesseract = None  # type: ignore

try:  # PDF image extraction
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover - optional
    fitz = None  # type: ignore
import zipfile


class DocumentIngestionService:
    def __init__(self):
        self.embedding = GeminiEmbeddings(
            api_key=settings.api_key,
            model=settings.gemini_embedding_model,
        )
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )

    def _load_pdf(self, file_bytes: bytes) -> str:
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            try:
                texts.append(page.extract_text() or "")
            except Exception:
                continue
        base_text = "\n".join(texts)

        # Attempt to extract text from images (figures/graphs) via OCR
        ocr_text = self._extract_pdf_images_text(file_bytes)
        return base_text + ("\n\n" + ocr_text if ocr_text else "")

    def _load_docx(self, file_bytes: bytes) -> str:
        doc = docx.Document(io.BytesIO(file_bytes))
        text_part = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        # Extract embedded images from DOCX zip and OCR
        ocr_text = self._extract_docx_images_text(file_bytes)
        return text_part + ("\n\n" + ocr_text if ocr_text else "")

    def _load_csv(self, file_bytes: bytes) -> str:
        df = pd.read_csv(io.BytesIO(file_bytes))
        return df.to_csv(index=False)

    def parse_file(self, filename: str, file_bytes: bytes) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(file_bytes)
        if suffix == ".docx":
            return self._load_docx(file_bytes)
        if suffix == ".csv":
            return self._load_csv(file_bytes)
        raise ValueError(f"Unsupported file type: {suffix}")

    def chunk(self, raw_text: str) -> List[Document]:
        return [Document(page_content=chunk) for chunk in self.splitter.split_text(raw_text)]

    def ingest(self, filename: str, file_bytes: bytes) -> int:
        """Parse, chunk, embed and upsert into Qdrant collection. Returns number of chunks."""
        raw_text = self.parse_file(filename, file_bytes)
        docs = self.chunk(raw_text)
        # Initialize or extend vector store
        qdrant_url = f"http://{settings.qdrant_host}:{settings.qdrant_port}"
        Qdrant.from_documents(
            docs,
            self.embedding,
            url=qdrant_url,
            collection_name=settings.documents_collection,
        )
        return len(docs)

    # ---------- Private helpers for OCR pipelines ----------
    def _ocr_image_bytes(self, img_bytes: bytes) -> str:
        """Run OCR on image bytes; return extracted text or empty string if not available.

        Requires Pillow and pytesseract. If pytesseract runtime is missing (no Tesseract installed),
        logs a debug message and returns empty string.
        """
        if Image is None or pytesseract is None:
            logger.debug("OCR skipped: Pillow or pytesseract not available")
            return ""
        try:
            with Image.open(io.BytesIO(img_bytes)) as im:
                # Convert to RGB to avoid mode issues
                if im.mode not in ("RGB", "L"):
                    im = im.convert("RGB")
                text = pytesseract.image_to_string(im)
                return text.strip()
        except Exception as e:
            logger.debug(f"OCR failed for an image: {e}")
            return ""

    def _extract_pdf_images_text(self, file_bytes: bytes) -> str:
        """Extract images from PDF via PyMuPDF and OCR them. Returns concatenated text.

        If PyMuPDF or OCR is unavailable, returns empty string.
        """
        if fitz is None:
            return ""
        texts: List[str] = []
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page_index in range(len(doc)):
                page = doc[page_index]
                for img in page.get_images(full=True):
                    try:
                        xref = img[0]
                        img_dict = doc.extract_image(xref)
                        img_bytes = img_dict.get("image")
                        if not img_bytes:
                            continue
                        t = self._ocr_image_bytes(img_bytes)
                        if t:
                            texts.append(f"[Page {page_index+1} image OCR]\n{t}")
                    except Exception:
                        continue
            doc.close()
        except Exception as e:
            logger.debug(f"PDF image extraction failed: {e}")
        return "\n\n".join(texts)

    def _extract_docx_images_text(self, file_bytes: bytes) -> str:
        """Extract embedded images from DOCX (zip) and OCR them. Returns concatenated text.

        DOCX stores images under word/media/*. We scan those entries and OCR.
        """
        texts: List[str] = []
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/") and not name.endswith("/"):
                        try:
                            data = zf.read(name)
                            t = self._ocr_image_bytes(data)
                            if t:
                                texts.append(f"[DOCX image OCR]\n{t}")
                        except Exception:
                            continue
        except Exception as e:
            logger.debug(f"DOCX image extraction failed: {e}")
        return "\n\n".join(texts)
